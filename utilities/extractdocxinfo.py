import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from zipfile import ZipFile
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def extract_docx_info(dfile):

    document = docx.Document(dfile)

    #extract text
    text = ''
    if isinstance(document, Document):
        parent_elm = document.element.body
    elif isinstance(document, _Cell):
        parent_elm = document._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            text = text + "\n" + Paragraph(child, document).text
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = text + "\n" +paragraph.text

    
    
    
    #extract fonts
    fonts = []
    for style in document.styles:
        try: 
            if style.font != None and style.font.name != None:
                if style.font.name not in fonts:
                    fonts.append(style.font.name)
        except:
            pass

    
    #extract n_tables
    n_tables = len(document.tables)


    #extract linkedin link and email(if present)
    linkedin = ''
    email = ''
    rels = document.part.rels
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            if rels[rel]._target.startswith('http://www.linkedin.com') and linkedin == '':
                linkedin = rels[rel]._target
            if rels[rel]._target.startswith('mailto') and email == '':
                email = rels[rel]._target[len('mailto')+1:]



    

    #extract n_images
    n_images = 0
    document = ZipFile(dfile)
    for name in document.namelist():
        if name.startswith('word/media/image'):
            n_images += 1
    
    return {"linkedin": linkedin, "n_tables": n_tables, "fonts": fonts, "n_images": n_images, "text": text, "email": email}


#file = open('../../../resume-test/Resume_1.docx', 'rb')
#print(extract_docx_info(file))
